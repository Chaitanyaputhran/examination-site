#!/usr/bin/env python3
"""
Generate Lab Report Word Document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading_with_style(doc, text, level=1, bold=True, size=14, align='center'):
    """Add a styled heading"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if align == 'center':
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph_with_style(doc, text, bold=False, size=12, align='left', spacing_before=0, spacing_after=0):
    """Add a styled paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'

    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    para.paragraph_format.space_before = Pt(spacing_before)
    para.paragraph_format.space_after = Pt(spacing_after)

    return para

def create_lab_report():
    """Create the lab report document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header - University Name
    add_paragraph_with_style(doc, "Birla Institute of Technology & Science, Pilani",
                            bold=True, size=16, align='center', spacing_after=6)

    add_paragraph_with_style(doc, "Work Integrated Learning Programme Division",
                            bold=True, size=14, align='center', spacing_after=12)

    # Add a line
    doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Semester and Course Info
    add_paragraph_with_style(doc, "Second Semester 2025-2026",
                            bold=True, size=13, align='center', spacing_before=12, spacing_after=6)

    add_paragraph_with_style(doc, "Full Stack Application Development (S2-25_SESAPZG503)",
                            bold=True, size=13, align='center', spacing_after=12)

    # Lab Sheet Title
    add_paragraph_with_style(doc, "Lab Sheet 1:",
                            bold=True, size=14, align='center', spacing_before=12, spacing_after=6)

    add_paragraph_with_style(doc, "Managing Source Code using Distributed Version Control (Git)",
                            bold=True, size=13, align='center', spacing_after=18)

    # Add some space
    doc.add_paragraph()

    # Project Title Section
    add_paragraph_with_style(doc, "Project Topic:",
                            bold=True, size=13, align='left', spacing_before=12, spacing_after=6)

    add_paragraph_with_style(doc, "Examination Portal - Online Examination Management System",
                            bold=False, size=12, align='left', spacing_after=12)

    # GitHub Repository
    add_paragraph_with_style(doc, "GitHub Repository:",
                            bold=True, size=13, align='left', spacing_before=6, spacing_after=6)

    github_para = doc.add_paragraph()
    github_run = github_para.add_run("https://github.com/Chaitanyaputhran/examination-site")
    github_run.font.size = Pt(11)
    github_run.font.name = 'Times New Roman'
    github_run.font.color.rgb = RGBColor(0, 0, 255)
    github_run.underline = True
    github_para.paragraph_format.space_after = Pt(18)

    # Project Description
    add_paragraph_with_style(doc, "Project Description:",
                            bold=True, size=13, align='left', spacing_before=6, spacing_after=6)

    description = (
        "A comprehensive full-stack online examination management system built with Spring Boot (Backend) "
        "and React (Frontend). The application features role-based access control for administrators and students, "
        "real-time exam taking with timer, unlimited test retakes, automated grading, performance analytics, "
        "and PDF report generation with email delivery."
    )
    add_paragraph_with_style(doc, description,
                            bold=False, size=11, align='justify', spacing_after=12)

    # Technology Stack
    add_paragraph_with_style(doc, "Technology Stack:",
                            bold=True, size=13, align='left', spacing_before=6, spacing_after=6)

    tech_para = doc.add_paragraph(style='List Bullet')
    tech_para.add_run("Backend: Java 17, Spring Boot 3.2.0, Spring Security, MySQL 8.0, JWT Authentication").font.size = Pt(11)

    tech_para2 = doc.add_paragraph(style='List Bullet')
    tech_para2.add_run("Frontend: React 18.2.0, Tailwind CSS 3.3.6, Vite 5.0.8, Axios, React Router").font.size = Pt(11)
    tech_para2.paragraph_format.space_after = Pt(18)

    # Add space before submission section
    doc.add_paragraph()
    doc.add_paragraph()

    # Add a line
    doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Submitted By Section
    add_paragraph_with_style(doc, "Submitted By:",
                            bold=True, size=13, align='left', spacing_before=12, spacing_after=12)

    # Student 1
    student1 = doc.add_paragraph()
    student1.add_run("Name: ").bold = True
    student1.add_run("Sameer Saxena")
    student1_run = student1.runs[0]
    student1_run.font.size = Pt(12)
    student1_run.font.name = 'Times New Roman'
    student1.runs[1].font.size = Pt(12)
    student1.runs[1].font.name = 'Times New Roman'
    student1.paragraph_format.space_after = Pt(6)

    roll1 = doc.add_paragraph()
    roll1.add_run("Roll No.: ").bold = True
    roll1.add_run("2025SL930009")
    roll1.runs[0].font.size = Pt(12)
    roll1.runs[0].font.name = 'Times New Roman'
    roll1.runs[1].font.size = Pt(12)
    roll1.runs[1].font.name = 'Times New Roman'
    roll1.paragraph_format.space_after = Pt(12)

    # Student 2
    student2 = doc.add_paragraph()
    student2.add_run("Name: ").bold = True
    student2.add_run("Chaitanya")
    student2.runs[0].font.size = Pt(12)
    student2.runs[0].font.name = 'Times New Roman'
    student2.runs[1].font.size = Pt(12)
    student2.runs[1].font.name = 'Times New Roman'
    student2.paragraph_format.space_after = Pt(6)

    roll2 = doc.add_paragraph()
    roll2.add_run("Roll No.: ").bold = True
    roll2.add_run("2025SL93013")
    roll2.runs[0].font.size = Pt(12)
    roll2.runs[0].font.name = 'Times New Roman'
    roll2.runs[1].font.size = Pt(12)
    roll2.runs[1].font.name = 'Times New Roman'

    # Save document
    filename = "Lab_Report_Examination_Portal.docx"
    doc.save(filename)
    print(f"✅ Lab report created successfully: {filename}")
    return filename

if __name__ == "__main__":
    create_lab_report()
